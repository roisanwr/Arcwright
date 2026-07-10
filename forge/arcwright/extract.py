"""
Multi-format extraction module.
Supports: PDF (digital + OCR), EPUB, MOBI, DOCX, TXT, HTML.

Each format is converted to clean markdown text for downstream chunking.
"""

import os
import sys
import time
import re
from pathlib import Path
from typing import Optional


# ─── Format Registry ─────────────────────────────────────────────

SUPPORTED_FORMATS = {
    ".pdf": "PDF Document (digital + scanned via OCR)",
    ".epub": "EPUB eBook",
    ".mobi": "MOBI eBook (Kindle)",
    ".azw3": "MOBI eBook (Kindle)",
    ".docx": "Word Document",
    ".doc": "Word Document",
    ".txt": "Plain Text",
    ".html": "HTML Document",
    ".htm": "HTML Document",
}

MARKUP_FORMATS = {".html", ".htm"}
ARCHIVE_FORMATS = {".epub", ".mobi", ".azw3"}


# ─── Main Entry Point ────────────────────────────────────────────

def extract_file(file_path: str, force_ocr: bool = True) -> str:
    """
    Detect file format and extract text to markdown.

    Args:
        file_path: Path to the file (PDF, EPUB, MOBI, DOCX, TXT, HTML)
        force_ocr: Only used for PDF — force OCR for scanned/image PDFs

    Returns:
        Full markdown text of the extracted content

    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path, force_ocr)
    elif ext == ".epub":
        return extract_epub(file_path)
    elif ext in (".mobi", ".azw3"):
        return extract_mobi(file_path)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_path)
    elif ext == ".txt":
        return extract_txt(file_path)
    elif ext in (".html", ".htm"):
        return extract_html(file_path)
    else:
        raise ValueError(
            f"Unsupported format '{ext}'. "
            f"Supported: {', '.join(SUPPORTED_FORMATS)}"
        )


# ─── PDF (Existing — marker-pdf) ─────────────────────────────────

def get_marker_path():
    """Find marker-pdf in the venv site-packages."""
    venv_base = os.path.expanduser("~/Arcwright/venv")
    lib = os.path.join(venv_base, "lib")
    if os.path.isdir(lib):
        for d in os.listdir(lib):
            site_pkg = os.path.join(lib, d, "site-packages")
            if os.path.isdir(site_pkg):
                return site_pkg
    return None


def extract_pdf(pdf_path: str, force_ocr: bool = True) -> str:
    """
    Extract PDF to markdown text using marker-pdf.
    Handles both digital (text-based) and scanned (image-based) PDFs.

    Args:
        pdf_path: Path to the PDF file
        force_ocr: If True, always use OCR (needed for scanned/image PDFs).
                   If False, try text extraction first, fall back to OCR.

    Returns:
        Full markdown text of the extracted PDF
    """
    marker_path = get_marker_path()
    if marker_path and marker_path not in sys.path:
        sys.path.insert(0, marker_path)

    from marker.converters.pdf import PdfConverter
    from marker.models import create_model_dict
    from marker.config.parser import ConfigParser
    from marker.output import text_from_rendered

    config = {
        "output_format": "markdown",
        "force_ocr": force_ocr,
        "disable_tqdm": True,
        "use_llm": False,
    }

    config_parser = ConfigParser(config)
    converter = PdfConverter(
        config=config_parser.generate_config_dict(),
        artifact_dict=create_model_dict(),
        processor_list=config_parser.get_processors(),
        renderer=config_parser.get_renderer(),
    )

    print(f"  📄 Extracting PDF: {os.path.basename(pdf_path)} (OCR={'ON' if force_ocr else 'AUTO'})...")
    start = time.time()
    rendered = converter(pdf_path)
    full_text, _, images = text_from_rendered(rendered)
    elapsed = time.time() - start

    print(f"  ✅ PDF done in {elapsed:.1f}s — {len(full_text):,} chars, {len(images)} images")
    return full_text


# ─── EPUB ─────────────────────────────────────────────────────────

def extract_epub(epub_path: str) -> str:
    """
    Extract EPUB to markdown text.
    Uses ebooklib to read EPUB structure + BeautifulSoup + markdownify.
    """
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    import markdownify

    print(f"  📖 Extracting EPUB: {os.path.basename(epub_path)}...")
    start = time.time()

    book = epub.read_epub(epub_path)
    chapters = []

    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            # Parse HTML content
            soup = BeautifulSoup(item.get_content(), "html.parser")

            # Remove script/style elements
            for tag in soup(["script", "style", "nav"]):
                tag.decompose()

            # Get the title from the document
            title = None
            h1 = soup.find("h1")
            h2 = soup.find("h2")
            if h1:
                title = h1.get_text(strip=True)
            elif h2:
                title = h2.get_text(strip=True)

            # Convert HTML to markdown
            html_content = str(soup)
            markdown_text = markdownify.markdownify(
                html_content,
                heading_style="ATX",
                strip=["img", "figure", "script", "style"],
            )

            # Clean up excessive whitespace
            markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
            markdown_text = markdown_text.strip()

            if markdown_text:
                if title and not markdown_text.startswith("#"):
                    markdown_text = f"# {title}\n\n{markdown_text}"
                chapters.append(markdown_text)

    full_text = "\n\n---\n\n".join(chapters)
    elapsed = time.time() - start

    print(f"  ✅ EPUB done in {elapsed:.1f}s — {len(full_text):,} chars, {len(chapters)} chapters")
    return full_text


# ─── MOBI / AZW3 (Kindle) ────────────────────────────────────────

def extract_mobi(mobi_path: str) -> str:
    """
    Extract MOBI/AZW3 to markdown text.
    Uses the `mobi` Python package to extract raw HTML, then markdownify.
    """
    import mobi
    import markdownify
    from bs4 import BeautifulSoup

    print(f"  📱 Extracting MOBI: {os.path.basename(mobi_path)}...")
    start = time.time()

    # mobi.extract returns (temp_dir, file_path) or raises
    temp_dir, file_path = mobi.extract(mobi_path)

    try:
        if file_path and os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

            # Check if it's HTML
            if "<html" in content.lower() or "<body" in content.lower():
                soup = BeautifulSoup(content, "html.parser")
                for tag in soup(["script", "style", "nav"]):
                    tag.decompose()

                markdown_text = markdownify.markdownify(
                    str(soup),
                    heading_style="ATX",
                    strip=["img", "figure"],
                )
                markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
            else:
                # Already plain text
                markdown_text = content.strip()
        else:
            markdown_text = ""
    finally:
        # Clean up temp directory
        import shutil
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)

    elapsed = time.time() - start
    print(f"  ✅ MOBI done in {elapsed:.1f}s — {len(markdown_text):,} chars")
    return markdown_text


# ─── DOCX ─────────────────────────────────────────────────────────

def extract_docx(docx_path: str) -> str:
    """
    Extract Word document (.docx) to markdown text.
    Uses python-docx to read paragraphs and tables.
    """
    from docx import Document

    print(f"  📝 Extracting DOCX: {os.path.basename(docx_path)}...")
    start = time.time()

    doc = Document(docx_path)
    sections = []
    current_heading_level = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""

        if "heading 1" in style_name or style_name.startswith("heading1"):
            sections.append(f"# {text}")
        elif "heading 2" in style_name or style_name.startswith("heading2"):
            sections.append(f"## {text}")
        elif "heading 3" in style_name or style_name.startswith("heading3"):
            sections.append(f"### {text}")
        else:
            # Check if it's a list item
            if para.style and "list" in style_name:
                sections.append(f"- {text}")
            else:
                sections.append(text)

    full_text = "\n\n".join(sections)
    elapsed = time.time() - start

    print(f"  ✅ DOCX done in {elapsed:.1f}s — {len(full_text):,} chars, {len(doc.paragraphs)} paragraphs")
    return full_text


# ─── TXT ──────────────────────────────────────────────────────────

def extract_txt(txt_path: str) -> str:
    """Extract plain text file as-is."""
    print(f"  📃 Extracting TXT: {os.path.basename(txt_path)}...")
    start = time.time()

    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    elapsed = time.time() - start
    print(f"  ✅ TXT done in {elapsed:.1f}s — {len(text):,} chars")
    return text


# ─── HTML ─────────────────────────────────────────────────────────

def extract_html(html_path: str) -> str:
    """
    Extract HTML to markdown text.
    Uses BeautifulSoup + markdownify.
    """
    import markdownify
    from bs4 import BeautifulSoup as BS

    print(f"  🌐 Extracting HTML: {os.path.basename(html_path)}...")
    start = time.time()

    with open(html_path, "r", encoding="utf-8", errors="replace") as f:
        html_content = f.read()

    soup = BS(html_content, "html.parser")

    # Remove non-content elements
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Try to get main content first
    main = soup.find("main") or soup.find("article") or soup.find("body") or soup

    markdown_text = markdownify.markdownify(
        str(main),
        heading_style="ATX",
        strip=["img", "figure", "script", "style", "nav"],
    )

    # Clean up
    markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
    markdown_text = re.sub(r'\[.*?\]\(.*?\)', '', markdown_text)  # Remove image links
    markdown_text = markdown_text.strip()

    elapsed = time.time() - start
    print(f"  ✅ HTML done in {elapsed:.1f}s — {len(markdown_text):,} chars")
    return markdown_text


# ─── Format Detection Helper ──────────────────────────────────────

def detect_format(file_path: str) -> str:
    """Return human-readable format name for a file."""
    ext = Path(file_path).suffix.lower()
    return SUPPORTED_FORMATS.get(ext, "Unknown format")
